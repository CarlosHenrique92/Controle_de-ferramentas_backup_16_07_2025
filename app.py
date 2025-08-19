from flask import Flask, render_template, request, redirect, url_for, flash, session
from datetime import datetime
import sqlite3
import openpyxl
import os
import smtplib
from email.message import EmailMessage
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from fpdf import FPDF
import unicodedata
from flask import send_file
import re
from docx import Document
from werkzeug.utils import secure_filename
import unicodedata, re

app = Flask(__name__)
app.secret_key = 'chave_secreta_segura' #Nescesário para uso da sessão

def ensure_schema():
    conn = sqlite3.connect('ferramentas.db')
    cur = conn.cursor()

    # --- garantir coluna perfuratriz em ferramentas ---
    cur.execute("PRAGMA table_info(ferramentas)")
    cols_ferr = [row[1] for row in cur.fetchall()]
    if 'perfuratriz' not in cols_ferr:
        cur.execute("ALTER TABLE ferramentas ADD COLUMN perfuratriz TEXT DEFAULT ''")

    # --- opcional: garantir modalidade_envio em requisicoes (se você usa isso no INSERT) ---
    cur.execute("PRAGMA table_info(requisicoes)")
    cols_req = [row[1] for row in cur.fetchall()]
    if 'modalidade_envio' not in cols_req:
        cur.execute("ALTER TABLE requisicoes ADD COLUMN modalidade_envio TEXT DEFAULT ''")

    conn.commit()
    conn.close()

# chame a migração ao iniciar
ensure_schema()

os.makedirs("uploads", exist_ok=True)

def _to_text(cell):
    # Junta todo texto de uma célula (às vezes vem com quebras)
    return " ".join(cell.text.split()).strip()

from docx import Document
import unicodedata, re

from docx import Document
import unicodedata, re

def _norm(txt: str) -> str:
    txt = txt or ""
    txt = unicodedata.normalize("NFKD", txt)
    txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
    return txt.strip().lower()

def _clean(txt: str) -> str:
    return " ".join((txt or "").split()).strip()

def _int_from(s: str) -> int:
    m = re.search(r"\d+", s or "")
    return int(m.group()) if m else 0

def _norm(txt: str) -> str:
    txt = txt or ""
    txt = unicodedata.normalize("NFKD", txt)
    txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
    return txt.strip().lower()

def _clean(txt: str) -> str:
    return " ".join((txt or "").split()).strip()

def _int_from(s: str) -> int:
    m = re.search(r"\d+", s or "")
    return int(m.group()) if m else 0

def parse_requisicao_docx(caminho):
    """
    Lê o DOCX com:
      Cabeçalho: Nome / Projeto / IDGEO / Perfuratriz (ou Placa)
      Tabela: ... | Quant./QTD/Quantidade | Descrição (ou Nome)
    Retorna técnico, local, idgeo, perfuratriz e itens (nome, quantidade).
    """
    doc = Document(caminho)
    dados = {
        "tecnico": "",
        "responsavel": "",
        "local": "",
        "idgeo": "",
        "perfuratriz": "",
        "ferramentas": []
    }

    # -------- 1) Cabeçalho vindo de PARÁGRAFOS (se houver) --------
    for p in doc.paragraphs:
        t = p.text.strip()
        if ":" not in t:
            continue
        k, v = t.split(":", 1)
        k_norm = _norm(k)
        v = v.strip()

        if k_norm == "nome":
            dados["tecnico"] = v
            dados["responsavel"] = v
        elif k_norm == "projeto":
            dados["local"] = v
        elif k_norm.startswith("idgeo"):
            dados["idgeo"] = v
        elif k_norm == "perfuratriz":
            dados["perfuratriz"] = v
        elif k_norm == "placa" and not dados["perfuratriz"]:
            dados["perfuratriz"] = v

    # -------- 2) Cabeçalho vindo de CÉLULAS DE TABELA (ex.: "Chave: Valor") --------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_txt = _clean(cell.text)
                if ":" not in cell_txt:
                    continue
                k, v = cell_txt.split(":", 1)
                k_norm = _norm(k)
                v = v.strip()

                if k_norm == "nome":
                    dados["tecnico"] = v
                    dados["responsavel"] = v
                elif k_norm == "projeto":
                    dados["local"] = v
                elif k_norm.startswith("idgeo"):
                    dados["idgeo"] = v
                elif k_norm == "perfuratriz":
                    dados["perfuratriz"] = v
                elif k_norm == "placa" and not dados["perfuratriz"]:
                    dados["perfuratriz"] = v

    # -------- 3) Itens: detectar a LINHA do cabeçalho da tabela --------
    qty_headers = {"quant", "quant.", "qtd", "quantidade"}
    desc_headers = {
        "descricao", "descrição", "descricao do item", "descrição do item",
        "nome", "nome do item"
    }

    for table in doc.tables:
        header_row_idx = None
        qty_idx = None
        desc_idx = None

        # Procura a linha que contém as colunas de quantidade e descrição
        for r_i, row in enumerate(table.rows):
            heads = [_norm(_clean(c.text)) for c in row.cells]
            if not heads:
                continue
            # encontra índices
            for i, h in enumerate(heads):
                if not h:
                    continue
                token = h.split()[0].strip(".:;")
                if qty_idx is None and token in qty_headers:
                    qty_idx = i
                if desc_idx is None and (h in desc_headers or h.startswith("descr")):
                    desc_idx = i
            if qty_idx is not None and desc_idx is not None:
                header_row_idx = r_i
                break

        if header_row_idx is None:
            continue  # essa tabela não é a de itens

        # Lê as linhas abaixo do cabeçalho
        for row in table.rows[header_row_idx + 1:]:
            cells = row.cells
            # ignora linhas/células mescladas e incompletas
            if len(cells) <= max(qty_idx, desc_idx):
                continue

            qtd_txt = _clean(cells[qty_idx].text)
            nome_txt = _clean(cells[desc_idx].text)

            if not nome_txt:
                continue

            qtd = _int_from(qtd_txt)
            if qtd <= 0:
                continue

            dados["ferramentas"].append({"nome": nome_txt, "quantidade": qtd})

    # trims finais
    for k in ("tecnico", "responsavel", "local", "idgeo", "perfuratriz"):
        dados[k] = (dados.get(k) or "").strip()

    return dados


def gerar_pdf_solicitacao(dados, ferramentas, caminho_pdf):
    pdf = FPDF()
    pdf.add_page()

    # === Inserir LOGO centralizada ===
    try:
        pdf.image("static/logo.png", x=80, w=50)  # Ajuste o caminho se a logo estiver em outro lugar
        pdf.ln(25)
    except:
        pass  # Caso a logo não esteja presente, o código continua normalmente

    # === TÍTULO CENTRALIZADO ===
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "REQUISIÇÃO DE FERRAMENTAS", ln=True, align='C')
    pdf.ln(10)

    # === DADOS DO CABEÇALHO ===
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 8, f" Requisição: {dados['nome_requisicao']}", ln=True)
    pdf.cell(0, 8, f" Responsável: {dados['responsavel']}", ln=True)
    pdf.cell(0, 8, f" Projeto/Local: {dados['local']}", ln=True)
    pdf.cell(0, 8, f" Técnico: {dados['tecnico']}", ln=True)
    pdf.cell(0, 8, f" IDGEO: {dados['idgeo']}", ln=True)
    pdf.cell(0, 8, f" Modalidade de Envio: {dados.get('modalidade_envio', 'Não especificada')}", ln=True)
    pdf.cell(0, 8, f" Data de Envio: {dados['data_envio']}", ln=True)
    pdf.cell(0, 8, f" Data da Solicitação: {dados['data_solicitacao']}", ln=True)

    pdf.ln(10)

    # === LISTA DE FERRAMENTAS (com borda e colunas) ===
    pdf.set_font("Arial", 'B', 12)
    pdf.set_fill_color(220, 220, 220)
    pdf.cell(100, 10, "Nome da Ferramenta", border=1, fill=True)
    pdf.cell(40, 10, "Quantidade UN", border=1, fill=True)

    pdf.ln()

    pdf.set_font("Arial", '', 12)
    for f in ferramentas:
        pdf.cell(100, 10, f['nome'], border=1)
        pdf.cell(40, 10, str(f['quantidade']), border=1)
        pdf.ln()

    pdf.output(caminho_pdf)


# -------------------- CONEXÃO --------------------
def get_db_connection():
    conn = sqlite3.connect('ferramentas.db')
    conn.row_factory = sqlite3.Row
    return conn

def enviar_email_com_anexo(dados, caminho_pdf):
    import smtplib
    from email.message import EmailMessage

    msg = EmailMessage()
    msg['Subject'] = f"Nova Requisição: {dados['nome_requisicao']}"
    msg['From'] = 'almoxarifado@geoambiente.eng.br'
    msg['To'] = 'almoxarifado@geoambiente.eng.br'

    corpo = f"""
Nova solicitação de ferramentas!

📄 Requisição: {dados['nome_requisicao']}
📦 Data de envio: {dados['data_envio']}
👤 Responsável: {dados['responsavel']}
🏗️ Local/Projeto: {dados['local']}
🔧 Técnico: {dados['tecnico']}
🆔 IDGEO: {dados['idgeo']}
🚚 Modalidade de envio: {dados.get('modalidade_envio', 'não especificada')}
📅 Solicitado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}

🔩 Ferramentas:
"""
    for f in dados['ferramentas']:
        corpo += f"- {f['nome']}: {f['quantidade']}\n"

    msg.set_content(corpo)

    with open(caminho_pdf, 'rb') as f:
        msg.add_attachment(f.read(), maintype='application', subtype='pdf', filename=os.path.basename(caminho_pdf))

    try:
        with smtplib.SMTP('smtp.gmail.com', 587) as smtp:
            smtp.starttls()
            smtp.login('almoxarifado@geoambiente.eng.br', 'swkb kwmk tjjm ozio')  # senha de app
            smtp.send_message(msg)
            print("📨 E-mail enviado com sucesso!")
    except Exception as e:
        print("❌ Erro ao enviar e-mail:", e)


# Cria a tabela 'requisicoes' caso ainda não exista
def criar_tabela_requisicoes():
    conn = sqlite3.connect('ferramentas.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS requisicoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome_requisicao TEXT,
            data_solicitacao TEXT,
            data_envio TEXT,
            responsavel TEXT,
            local TEXT,
            tecnico TEXT,
            idgeo TEXT,
            ferramentas TEXT
        )
    ''')
    conn.commit()
    conn.close()

# Chama a função para garantir que a tabela existe
criar_tabela_requisicoes()




# -------------------- LOGIN --------------------
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        senha = request.form['senha']
        if senha == 'Geo@#07981':
            session['logado'] = True
            return redirect('/')
        else:
            return render_template('login.html', erro='Senha incorreta.')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/')

# -------------------- PÁGINA INICIAL --------------------
@app.route('/')
def index():
    try:
        conn = get_db_connection()
        ferramentas_estoque = conn.execute('SELECT * FROM ferramentas WHERE status = "estoque"').fetchall()
        ferramentas_uso = conn.execute('SELECT * FROM ferramentas WHERE status = "uso"').fetchall()
        conn.close()
        return render_template('index.html', ferramentas_estoque=ferramentas_estoque, ferramentas_uso=ferramentas_uso)
    except sqlite3.OperationalError as e:
        erro = str(e)
        return render_template('index.html', ferramentas_estoque=[], ferramentas_uso=[], erro=erro)

# -------------------- ADICIONAR --------------------
@app.route('/adicionar', methods=['POST'])
def adicionar():
    if not session.get('logado'):
        return redirect('/login')

    nome = request.form['nome']
    status = request.form['status']
    local = request.form['local']
    tecnico = request.form['tecnico']
    quantidade = int(request.form['quantidade'])
    idgeo = request.form['idgeo']
    perfuratriz = request.form['perfuratriz']  # <- NOVO

    conn = get_db_connection()

    if status == 'estoque':
        existente = conn.execute('SELECT * FROM ferramentas WHERE nome = ? AND status = "estoque"', (nome,)).fetchone()
        if existente:
            nova_qtd = existente['quantidade'] + quantidade
            conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (nova_qtd, existente['id']))
        else:
            conn.execute('''INSERT INTO ferramentas (nome, status, local, tecnico, quantidade, idgeo, perfuratriz)
                            VALUES (?, ?, ?, ?, ?, ?, ?)''',
                            (nome, status, '', '', quantidade, '', ''))  # estoque não tem local nem técnico
    else:  # Em uso
        existente = conn.execute('''SELECT * FROM ferramentas
                                    WHERE nome = ? AND status = "uso" AND local = ? AND tecnico = ? AND idgeo = ? AND perfuratriz = ?''',
                                    (nome, local, tecnico, idgeo, perfuratriz)).fetchone()
        if existente:
            nova_qtd = existente['quantidade'] + quantidade
            conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (nova_qtd, existente['id']))
        else:
            conn.execute('''INSERT INTO ferramentas (nome, status, local, tecnico, quantidade, idgeo, perfuratriz)
                            VALUES (?, ?, ?, ?, ?, ?, ?)''',
                            (nome, status, local, tecnico, quantidade, idgeo, perfuratriz))

    conn.commit()
    conn.close()
    return redirect('/')


# -------------------- EDITAR --------------------
@app.route('/editar/<int:id>', methods=['GET', 'POST'])
def editar(id):
    if not session.get('logado'):
        return redirect('/login')

    conn = get_db_connection()
    ferramenta = conn.execute('SELECT * FROM ferramentas WHERE id = ?', (id,)).fetchone()

    if not ferramenta:
        conn.close()
        flash('Ferramenta não encontrada.', 'warning')
        return redirect('/')

    if request.method == 'POST':
        nome = ferramenta['nome']
        status_anterior = ferramenta['status']
        quantidade_antiga = int(ferramenta['quantidade'])

        novo_status = request.form.get('status')
        try:
            quantidade_nova = max(0, int(request.form.get('quantidade', 0)))
        except (TypeError, ValueError):
            quantidade_nova = 0

        local = (request.form.get('local') or '').strip()
        tecnico = (request.form.get('tecnico') or '').strip()
        idgeo = (request.form.get('idgeo') or '').strip()
        perfuratriz = (request.form.get('perfuratriz') or '').strip()

        # ========== ESTOQUE -> USO ==========
        if status_anterior == 'estoque' and novo_status == 'uso':
            # saldo no estoque após envio
            saldo_estoque = max(quantidade_antiga - quantidade_nova, 0)
            # atualiza o saldo no próprio registro de estoque
            conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (saldo_estoque, id))

            # em uso: perfuratriz é opcional; mesclar por (nome, local, tecnico, idgeo, perfuratriz ou vazio)
            if perfuratriz:
                existente_uso = conn.execute(
                    '''SELECT * FROM ferramentas
                       WHERE nome = ? AND status = 'uso'
                         AND local = ? AND tecnico = ? AND idgeo = ? AND perfuratriz = ?''',
                    (nome, local, tecnico, idgeo, perfuratriz)
                ).fetchone()
            else:
                existente_uso = conn.execute(
                    '''SELECT * FROM ferramentas
                       WHERE nome = ? AND status = 'uso'
                         AND local = ? AND tecnico = ? AND idgeo = ?
                         AND (perfuratriz IS NULL OR perfuratriz = '')''',
                    (nome, local, tecnico, idgeo)
                ).fetchone()

            if quantidade_nova > 0:
                if existente_uso:
                    nova_qtd_uso = int(existente_uso['quantidade']) + quantidade_nova
                    conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (nova_qtd_uso, existente_uso['id']))
                else:
                    conn.execute(
                        '''INSERT INTO ferramentas (nome, status, local, tecnico, quantidade, idgeo, perfuratriz)
                           VALUES (?, 'uso', ?, ?, ?, ?, ?)''',
                        (nome, local, tecnico, quantidade_nova, idgeo, perfuratriz)
                    )

        # ========== USO -> ESTOQUE ==========
        elif status_anterior == 'uso' and novo_status == 'estoque':
            existente_estoque = conn.execute(
                'SELECT * FROM ferramentas WHERE nome = ? AND status = "estoque"', (nome,)
            ).fetchone()

            if quantidade_nova > 0:
                if existente_estoque:
                    nova_qtd_est = int(existente_estoque['quantidade']) + quantidade_nova
                    conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (nova_qtd_est, existente_estoque['id']))
                else:
                    conn.execute(
                        '''INSERT INTO ferramentas (nome, status, local, tecnico, quantidade, idgeo, perfuratriz)
                           VALUES (?, 'estoque', '', '', ?, '', '')''',
                        (nome, quantidade_nova)
                    )

            # saldo que permanece em uso com o técnico
            quantidade_restante = quantidade_antiga - quantidade_nova
            if quantidade_restante > 0:
                conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (quantidade_restante, id))
            else:
                conn.execute('DELETE FROM ferramentas WHERE id = ?', (id,))

        # ========== USO -> USO (pode mudar local/técnico/idgeo/perfuratriz) ==========
        elif status_anterior == 'uso' and novo_status == 'uso':
            if perfuratriz:
                existente_uso = conn.execute(
                    '''SELECT * FROM ferramentas
                       WHERE nome = ? AND status = 'uso'
                         AND local = ? AND tecnico = ? AND idgeo = ? AND perfuratriz = ?''',
                    (nome, local, tecnico, idgeo, perfuratriz)
                ).fetchone()
            else:
                existente_uso = conn.execute(
                    '''SELECT * FROM ferramentas
                       WHERE nome = ? AND status = 'uso'
                         AND local = ? AND tecnico = ? AND idgeo = ?
                         AND (perfuratriz IS NULL OR perfuratriz = '')''',
                    (nome, local, tecnico, idgeo)
                ).fetchone()

            if existente_uso and existente_uso['id'] != id:
                # mescla: soma na linha existente e remove a atual
                nova_qtd = int(existente_uso['quantidade']) + quantidade_nova
                conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (nova_qtd, existente_uso['id']))
                conn.execute('DELETE FROM ferramentas WHERE id = ?', (id,))
            else:
                # atualiza a própria linha
                conn.execute(
                    '''UPDATE ferramentas
                       SET local = ?, tecnico = ?, idgeo = ?, perfuratriz = ?, quantidade = ?
                       WHERE id = ?''',
                    (local, tecnico, idgeo, perfuratriz, quantidade_nova, id)
                )

        # ========== ESTOQUE -> ESTOQUE (apenas ajustar quantidade) ==========
        else:
            # estoque não guarda local/tecnico/idgeo/perfuratriz
            conn.execute(
                '''UPDATE ferramentas
                   SET status = 'estoque', quantidade = ?, local = '', tecnico = '', idgeo = '', perfuratriz = ''
                   WHERE id = ?''',
                (quantidade_nova, id)
            )

        conn.commit()
        conn.close()
        flash('Ferramenta atualizada com sucesso!', 'success')
        return redirect('/')

    conn.close()
    return render_template('editar.html', ferramenta=ferramenta)



# -------------------- DELETAR --------------------
@app.route('/deletar/<int:id>')
def deletar(id):
    if not session.get('logado'):
        return redirect('/login')

    conn = get_db_connection()
    conn.execute('DELETE FROM ferramentas WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    return redirect('/')

# -------------------- DEVOLVER --------------------
@app.route('/devolver/<int:id>')
def devolver(id):
    if not session.get('logado'):
        return redirect('/login')

    conn = get_db_connection()
    ferramenta = conn.execute('SELECT * FROM ferramentas WHERE id = ?', (id,)).fetchone()
    if ferramenta:
        existente = conn.execute('SELECT * FROM ferramentas WHERE nome = ? AND status = "estoque"', (ferramenta['nome'],)).fetchone()
        if existente:
            nova_qtd = existente['quantidade'] + ferramenta['quantidade']
            conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (nova_qtd, existente['id']))
            conn.execute('DELETE FROM ferramentas WHERE id = ?', (ferramenta['id'],))
        else:
            conn.execute('UPDATE ferramentas SET status = "estoque", local = "", tecnico = "", idgeo = "" WHERE id = ?', (ferramenta['id'],))

    conn.commit()
    conn.close()
    return redirect('/')

# -------------------- RELATÓRIOS --------------------
@app.route('/relatorios', methods=['GET'])
def relatorios():
    conn = get_db_connection()
    ferramenta = request.args.get('ferramenta', '').lower()
    tecnico = request.args.get('tecnico', '').lower()
    projeto = request.args.get('projeto', '').lower()
    idgeo = request.args.get('idgeo', '').lower()

    query = '''SELECT nome, quantidade, status, local, tecnico, idgeo, perfuratriz FROM ferramentas WHERE status = 'uso' '''
    params = []
    if ferramenta:
        query += ' AND LOWER(nome) LIKE ?'
        params.append(f'%{ferramenta}%')
    if tecnico:
        query += ' AND LOWER(tecnico) LIKE ?'
        params.append(f'%{tecnico}%')
    if projeto:
        query += ' AND LOWER(local) LIKE ?'
        params.append(f'%{projeto}%')
    if idgeo:
        query += ' AND LOWER(idgeo) LIKE ?'
        params.append(f'%{idgeo}%')

    ferramentas = conn.execute(query, params).fetchall()
    nomes_ferramentas = [row['nome'] for row in conn.execute("SELECT DISTINCT nome FROM ferramentas WHERE status = 'uso'")]
    tecnicos = [row['tecnico'] for row in conn.execute("SELECT DISTINCT tecnico FROM ferramentas WHERE status = 'uso' AND tecnico != ''")]
    locais = [row['local'] for row in conn.execute("SELECT DISTINCT local FROM ferramentas WHERE status = 'uso' AND local != ''")]
    idgeos = [row['idgeo'] for row in conn.execute("SELECT DISTINCT idgeo FROM ferramentas WHERE status = 'uso' AND idgeo != ''")]

    conn.close()
    return render_template('relatorios.html', ferramentas=ferramentas, nomes_ferramentas=nomes_ferramentas, tecnicos=tecnicos, locais=locais, idgeos=idgeos)

# -------------------- RELATÓRIO ESTOQUE --------------------
@app.route('/relatorio_estoque', methods=['GET', 'POST'])
def relatorio_estoque():
    conn = get_db_connection()
    nome_filtro = ''
    if request.method == 'POST':
        nome_filtro = request.form.get('nome', '').lower()
        query = "SELECT nome, quantidade FROM ferramentas WHERE status = 'estoque'"
        params = []
        if nome_filtro:
            query += " AND LOWER(nome) LIKE ?"
            params.append(f"%{nome_filtro}%")
        ferramentas_estoque = conn.execute(query, params).fetchall()
    else:
        ferramentas_estoque = conn.execute("SELECT nome, quantidade FROM ferramentas WHERE status = 'estoque'").fetchall()

    nomes = [row['nome'] for row in conn.execute("SELECT DISTINCT nome FROM ferramentas WHERE status = 'estoque'")]
    conn.close()
    return render_template('relatorio_estoque.html', ferramentas_estoque=ferramentas_estoque, nomes=nomes)

# -------------------- EXPORTAÇÃO --------------------
@app.route('/exportar_excel')
def exportar_excel():
    ferramenta = request.args.get('ferramenta', '').lower()
    tecnico = request.args.get('tecnico', '').lower()
    projeto = request.args.get('projeto', '').lower()
    idgeo = request.args.get('idgeo', '').lower()

    # Consulta SQL com filtro e inclusão da coluna perfuratriz
    query = '''SELECT nome, quantidade, status, local, tecnico, idgeo, perfuratriz FROM ferramentas WHERE status = 'uso' '''
    params = []
    if ferramenta:
        query += ' AND LOWER(nome) LIKE ?'
        params.append(f'%{ferramenta}%')
    if tecnico:
        query += ' AND LOWER(tecnico) LIKE ?'
        params.append(f'%{tecnico}%')
    if projeto:
        query += ' AND LOWER(local) LIKE ?'
        params.append(f'%{projeto}%')
    if idgeo:
        query += ' AND LOWER(idgeo) LIKE ?'
        params.append(f'%{idgeo}%')

    # Conexão e execução
    conn = sqlite3.connect('ferramentas.db')
    cursor = conn.cursor()
    dados = cursor.execute(query, params).fetchall()
    conn.close()

    if not dados:
        return "Nenhum dado encontrado para exportar."

    # Geração do Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Ferramentas em Uso"

    # Cabeçalho
    ws.append(["Nome", "Quantidade", "Status", "Projeto", "Técnico", "IDGEO", "Perfuratriz"])

    # Dados
    for linha in dados:
        ws.append(linha)

    # Exportar
    nome_arquivo = "relatorio_projetos.xlsx"
    wb.save(nome_arquivo)
    return send_file(nome_arquivo, as_attachment=True)


@app.route('/exportar_estoque')
def exportar_estoque():
    conn = get_db_connection()
    ferramentas = conn.execute('SELECT * FROM ferramentas WHERE status = "estoque"').fetchall()
    conn.close()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Estoque'
    ws.append(['ID', 'Nome', 'Quantidade', 'Local', 'Técnico', 'IDGEO'])
    for f in ferramentas:
        ws.append([f['id'], f['nome'], f['quantidade'], f['local'], f['tecnico'], f['idgeo']])

    caminho_arquivo = os.path.join(os.path.expanduser("~"), "relatorio_estoque.xlsx")
    wb.save(caminho_arquivo)
    return send_file(caminho_arquivo, as_attachment=True)

#--------------CONFIRMAR RELATORIO------------------------------------
@app.route('/confirmar_solicitacao', methods=['POST'])
def confirmar_solicitacao():
    ferramentas_selecionadas = request.form.getlist('ferramentas[]')
    dados = {
        'nome_requisicao': request.form.get('nome_requisicao'),
        'responsavel': request.form.get('responsavel'),
        'local': request.form.get('local'),
        'tecnico': request.form.get('tecnico'),
        'idgeo': request.form.get('idgeo'),
        'data_envio': request.form.get('data_envio'),
        'modalidade_envio': request.form.get('modalidade_envio'),
        'ferramentas': []
    }

    for nome in ferramentas_selecionadas:
        qtd = request.form.get(f'quantidade_{nome}', '')
        dados['ferramentas'].append({
            'nome': nome,
            'quantidade': qtd
        })

    session['dados_requisicao'] = dados
    return render_template('confirmar_solicitacao.html', dados=dados)

# --- Rota final para processar a requisição ---
@app.route('/solicitar', methods=['GET', 'POST'])
def solicitar_ferramentas():
    conn = get_db_connection()

    if request.method == 'POST':
        # Se for retorno da tela de confirmação (botão Voltar)
        if request.form.get('confirmacao_final') == 'true':
            ferramentas_form = []
            ferramentas_disponiveis = conn.execute(
                'SELECT nome, quantidade FROM ferramentas WHERE status = "estoque" AND quantidade > 0'
            ).fetchall()

            ferramentas_selecionadas = request.form.getlist('ferramentas[]')

            for f in ferramentas_disponiveis:
                nome = f['nome']
                ferramentas_form.append({
                    'nome': nome,
                    'quantidade': f['quantidade'],
                    'selecionada': nome in ferramentas_selecionadas,
                    'qtd_solicitada': request.form.get(f'quantidade_' + nome, '')
                })

            conn.close()
            return render_template('solicitar.html', ferramentas=ferramentas_form,
                                   nome_requisicao=request.form.get('nome_requisicao'),
                                   responsavel=request.form.get('responsavel'),
                                   local=request.form.get('local'),
                                   tecnico=request.form.get('tecnico'),
                                   idgeo=request.form.get('idgeo'),
                                   data_envio=request.form.get('data_envio'),
                                   modalidade_envio=request.form.get('modalidade_envio'))

        # Processa a confirmação final e envia o e-mail
        if request.form.get('confirmado') == 'true':
            data_envio_raw = request.form.get('data_envio')  # exemplo: '2025-07-31'
            data_envio_formatada = datetime.strptime(data_envio_raw, '%Y-%m-%d').strftime('%d/%m/%Y')
            ferramentas_selecionadas = request.form.getlist('ferramentas[]')
            dados = {
                'nome_requisicao': request.form.get('nome_requisicao'),
                'responsavel': request.form.get('responsavel'),
                'local': request.form.get('local'),
                'tecnico': request.form.get('tecnico'),
                'idgeo': request.form.get('idgeo'),
                'data_envio':data_envio_formatada,
                'modalidade_envio': request.form.get('modalidade_envio'),
                'ferramentas': []
            }

            for nome in ferramentas_selecionadas:
                qtd = request.form.get(f'quantidade_{nome}', '1')
                dados['ferramentas'].append({'nome': nome, 'quantidade': qtd})

            data_solicitacao = datetime.now().strftime('%d/%m/%Y às %H:%M')
            dados['data_solicitacao'] = data_solicitacao
            ferramentas_str = ", ".join([f"{f['nome']} ({f['quantidade']})" for f in dados['ferramentas']])
            conn.execute('''
                INSERT INTO requisicoes (nome_requisicao, data_solicitacao, data_envio, responsavel, local, tecnico, idgeo, ferramentas, modalidade_envio)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (dados['nome_requisicao'], data_solicitacao, dados['data_envio'], dados['responsavel'],
                  dados['local'], dados['tecnico'], dados['idgeo'], ferramentas_str, dados['modalidade_envio']))

            # Atualizar estoque
            for f in dados['ferramentas']:
                nome = f['nome']
                quantidade = int(f['quantidade'])

                ferramenta_estoque = conn.execute(
                    'SELECT * FROM ferramentas WHERE nome = ? AND status = "estoque"', (nome,)
                ).fetchone()

                if ferramenta_estoque and ferramenta_estoque['quantidade'] >= quantidade:
                    nova_qtd = ferramenta_estoque['quantidade'] - quantidade
                    conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (nova_qtd, ferramenta_estoque['id']))

                    ferramenta_uso = conn.execute('''
                        SELECT * FROM ferramentas WHERE nome = ? AND status = "uso" AND local = ? AND tecnico = ? AND idgeo = ?
                    ''', (nome, dados['local'], dados['tecnico'], dados['idgeo'])).fetchone()

                    if ferramenta_uso:
                        nova_qtd_uso = ferramenta_uso['quantidade'] + quantidade
                        conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?', (nova_qtd_uso, ferramenta_uso['id']))
                    else:
                        conn.execute('''
                            INSERT INTO ferramentas (nome, status, local, tecnico, quantidade, idgeo)
                            VALUES (?, 'uso', ?, ?, ?, ?)
                        ''', (nome, dados['local'], dados['tecnico'], quantidade, dados['idgeo']))

            # Geração e envio do PDF
            nome_limpo = unicodedata.normalize('NFKD', dados['nome_requisicao']).encode('ASCII', 'ignore').decode('ASCII')
            nome_arquivo_pdf = f"requisicao_{nome_limpo.replace(' ', '_')}.pdf"
            pdf_path = os.path.join('static', nome_arquivo_pdf)
            gerar_pdf_solicitacao(dados, dados['ferramentas'], pdf_path)
            enviar_email_com_anexo(dados, pdf_path)

            conn.commit()
            conn.close()

            return render_template('sucesso.html', nome_pdf=nome_arquivo_pdf)

    # GET padrão: exibe o formulário
    ferramentas = conn.execute(
        'SELECT nome, quantidade FROM ferramentas WHERE status = "estoque" AND quantidade > 0'
    ).fetchall()
    conn.close()
    return render_template('solicitar.html', ferramentas=ferramentas)


@app.route('/sucesso')
def sucesso():
    nome_pdf = request.args.get('pdf')
    return render_template('sucesso.html', nome_pdf=nome_pdf)

#--------------------Devolução------------------------------
@app.route('/confirmar_devolucao', methods=['POST'])
def confirmar_devolucao():
    selecionadas = request.form.getlist('selecionadas')

    if not selecionadas:
        flash('Nenhuma ferramenta foi selecionada para devolução.', 'warning')
        return redirect(url_for('index'))

    conn = get_db_connection()
    ferramentas = conn.execute(
        f"SELECT * FROM ferramentas WHERE id IN ({','.join('?' * len(selecionadas))})",
        selecionadas
    ).fetchall()
    conn.close()

    return render_template('confirmar_devolucao.html', ferramentas=ferramentas)


@app.route('/executar_devolucao', methods=['POST'])
def executar_devolucao():
    selecionadas = request.form.getlist('selecionadas')

    if not selecionadas:
        flash('Nenhuma ferramenta foi selecionada para devolução.', 'warning')
        return redirect(url_for('index') + '#aba-uso')

    conn = get_db_connection()

    for ferramenta_id in selecionadas:
        ferramenta = conn.execute('SELECT * FROM ferramentas WHERE id = ?', (ferramenta_id,)).fetchone()

        if ferramenta:
            nome = ferramenta['nome']
            quantidade = ferramenta['quantidade']

            # Tenta encontrar outra ferramenta IGUAL no estoque
            existente = conn.execute('''
                SELECT * FROM ferramentas
                WHERE nome = ? AND status = 'estoque'
            ''', (nome,)).fetchone()

            if existente:
                nova_quantidade = existente['quantidade'] + quantidade

                # Atualiza o registro já existente no estoque
                conn.execute('''
                    UPDATE ferramentas
                    SET quantidade = ?
                    WHERE id = ?
                ''', (nova_quantidade, existente['id']))

                # Exclui o item atual (em uso), já que foi somado ao estoque
                conn.execute('DELETE FROM ferramentas WHERE id = ?', (ferramenta_id,))
            else:
                # Se não existe no estoque, apenas atualiza status e limpa técnico/local
                conn.execute('''
                    UPDATE ferramentas
                    SET status = 'estoque',
                        local = '',
                        tecnico = ''
                    WHERE id = ?
                ''', (ferramenta_id,))

    conn.commit()
    conn.close()

    flash('Ferramentas devolvidas com sucesso!', 'success')
    return redirect(url_for('index') + '#aba-uso')

#------------IMPORTAR CHECK LIST-------------------
@app.route('/uso/importar', methods=['GET', 'POST'])
def uso_importar():
    if not session.get('logado'):
        return redirect('/login')

    if request.method == 'POST':
        arq = request.files.get('arquivo')
        if not arq or not arq.filename.lower().endswith('.docx'):
            flash('Envie um arquivo .docx válido.', 'warning')
            return redirect(request.url)

        caminho = os.path.join('uploads', secure_filename(arq.filename))
        arq.save(caminho)

        try:
            dados = parse_requisicao_docx(caminho)  # <- função que já te passei
        except Exception as e:
            flash(f'Não foi possível ler o arquivo: {e}', 'danger')
            return redirect(request.url)

        tecnico = (dados.get('tecnico') or dados.get('responsavel') or '').strip()
        local = (dados.get('local') or '').strip()
        idgeo = (dados.get('idgeo') or '').strip()
        perfuratriz = (dados.get('perfuratriz') or '').strip()
        

        conn = get_db_connection()

        for f in dados.get('ferramentas', []):
            nome = f['nome'].strip()
            qtd_doc = int(f['quantidade'])

            # 1) Lançar/mesclar em USO (DOC sempre cria/atualiza)
            existente_uso = conn.execute(
                '''SELECT * FROM ferramentas
                   WHERE nome = ? AND status = "uso"
                     AND local = ? AND tecnico = ? AND idgeo = ? AND perfuratriz = ?''',
                (nome, local, tecnico, idgeo, perfuratriz)
            ).fetchone()

            if existente_uso:
                nova_qtd_uso = int(existente_uso['quantidade']) + qtd_doc
                conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?',
                             (nova_qtd_uso, existente_uso['id']))
            else:
                conn.execute(
                    '''INSERT INTO ferramentas (nome, status, local, tecnico, quantidade, idgeo, perfuratriz)
                       VALUES (?, "uso", ?, ?, ?, ?, ?)''',
                    (nome, local, tecnico, qtd_doc, idgeo, perfuratriz)
                )

            # 2) Abater/Zerar ESTOQUE se existir (físico > digital)
            estoque = conn.execute(
                'SELECT * FROM ferramentas WHERE nome = ? AND status = "estoque"', (nome,)
            ).fetchone()

            if estoque:
                qtd_estoque = int(estoque['quantidade'] or 0)
                if qtd_estoque <= 0:
                    # já está zerado, nada a fazer
                    pass
                else:
                    # abate até zerar (se DOC pedir mais do que há)
                    nova_qtd_est = max(qtd_estoque - qtd_doc, 0)
                    conn.execute('UPDATE ferramentas SET quantidade = ? WHERE id = ?',
                                 (nova_qtd_est, estoque['id']))

        conn.commit()
        conn.close()

        flash('Requisição importada: itens lançados em USO e estoque abatido/zerado quando aplicável.', 'success')
        return redirect(url_for('index') + '#aba-uso')

    # GET -> formulário simples
    return render_template('uso_importar.html')



if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')

#git add .
#git commit -m "Add importar formulario"
#git push

